import json
import io
from typing import IO, Any  # Imported IO for file types and Any for the style fix
from docx import Document
from docx.shared import Pt

# Added type hint: file is an IO object (like a file definition)
def json_to_str(file: IO[Any]) -> str:
    data = json.load(file)
    data_str = json.dumps(data, default=str, ensure_ascii=False)
    return data_str

def generate_docx_stream(agreement_text: str) -> io.BytesIO:
    """
    Generates a DOCX file in memory from the provided text string.
    Returns a BytesIO object (file stream) that can be sent directly to the frontend.
    """
    doc = Document()
    
    # Optional: Set a default style or font if needed
    # We hint 'style' as Any because Pylance thinks doc.styles['Normal'] 
    # is a generic 'BaseStyle' which doesn't have a .font attribute.
    style: Any = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Split the text by the form-feed character (\x0c) to handle explicit page breaks
    # or just split by newlines if your LLM doesn't output \x0c
    # The logic below handles both continuous text and explicit page breaks.
    
    pages = agreement_text.split('\x0c')
    
    for i, page_content in enumerate(pages):
        # Add paragraphs for each line to maintain basic formatting
        lines = page_content.split('\n')
        for line in lines:
            # Remove excessive whitespace but keep structure
            clean_line = line.strip() 
            if clean_line:
                doc.add_paragraph(clean_line)
            elif not clean_line:
                # Add empty paragraph for spacing
                doc.add_paragraph("")
        
        # Add a page break after each section, except the last one
        if i < len(pages) - 1:
            doc.add_page_break()

    # Save the document to an in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    
    # Reset the buffer position to the beginning so it can be read
    buffer.seek(0)
    
    return buffer